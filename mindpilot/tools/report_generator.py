"""
多格式报告生成器
================
输出格式：Word (.docx) + Markdown (.md) + HTML (.html)

报告结构（学术规范）：
  封面 / 摘要 / 目录
  一、研究背景与问题陈述
  二、文献综述
  三、实验设计与方法论
  四、核心代码实现
  五、实验结果与数据分析
  六、结论与展望
  参考文献
"""

import json
import re
import time
from pathlib import Path
from typing import Any, Optional
from datetime import datetime


class ReportGenerator:
    def __init__(self, output_dir: str = "outputs", logger=None):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.logger = logger

    # ── 主入口 ──────────────────────────────────────────────
    def generate(self, content: dict, filename: str = "report",
                 formats: list[str] = None) -> dict[str, str]:
        formats = formats or ["docx", "markdown", "html"]
        results = {}
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")

        for fmt in formats:
            try:
                if fmt == "docx":
                    path = self.output_dir / f"{filename}_{ts}.docx"
                    self._to_docx(content, path)
                elif fmt == "markdown":
                    text = self._to_markdown(content)
                    path = self.output_dir / f"{filename}_{ts}.md"
                    path.write_text(text, encoding="utf-8")
                elif fmt == "html":
                    text = self._to_html(content)
                    path = self.output_dir / f"{filename}_{ts}.html"
                    path.write_text(text, encoding="utf-8")
                else:
                    continue
                results[fmt] = str(path)
                if self.logger:
                    self.logger.success("ReportGen",
                        f"{fmt.upper()} 报告已保存: {path.name}")
            except Exception as e:
                if self.logger:
                    self.logger.error("ReportGen", f"{fmt} 生成失败: {e}")
                import traceback; traceback.print_exc()
        return results

    # ── Word 文档生成 ────────────────────────────────────────
    def _to_docx(self, c: dict, path: Path):
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # 页面设置 A4
        section = doc.sections[0]
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2.5)

        # 设置默认字体
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(12)

        def _add_heading(text, level=1):
            p = doc.add_heading(text, level=level)
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(13)
            else:
                run.font.size = Pt(12)
            return p

        def _add_body(text, bold=False, indent=False):
            if not text or not str(text).strip():
                return
            p = doc.add_paragraph()
            if indent:
                p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(str(text))
            run.font.size = Pt(12)
            run.bold = bold
            p.paragraph_format.space_after = Pt(6)
            return p

        def _add_code_block(code_text):
            if not code_text or not str(code_text).strip():
                return
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(str(code_text))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            # 灰色背景
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F2F2F2')
            pPr.append(shd)

        title  = c.get("title", "MindPilot 科研报告")
        query  = c.get("query", "")
        ts_str = datetime.now().strftime("%Y年%m月%d日 %H:%M")

        # ── 封面 ──
        doc.add_heading("", 0)   # 占位，清空默认样式
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(title)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.add_run(f"生成时间：{ts_str}  |  MindPilot 多 Agent 科研系统").font.size = Pt(11)
        doc.add_paragraph()

        # ── 摘要 ──
        _add_heading("摘　要", 1)
        abstract = c.get("abstract", "")
        if not abstract:
            abstract = (f"本报告由 MindPilot 多 Agent 科研助手系统自动生成，"
                        f"针对研究问题「{query}」，完成了文献检索、实验设计、"
                        f"代码实现、数据分析等全流程科研任务，并综合各模块输出形成本报告。")
        _add_body(abstract)
        doc.add_paragraph()

        # ── 各正文章节 ──
        sections = c.get("sections", [])
        for sec in sections:
            heading = sec.get("heading", "")
            body    = sec.get("body", "")
            level   = sec.get("level", 1)
            if heading:
                _add_heading(heading, level)
            if body:
                # 分段落输出
                for para in str(body).split("\n\n"):
                    para = para.strip()
                    if para:
                        _add_body(para)
            doc.add_paragraph()

        # ── 代码实现 ──
        code = c.get("code", "")
        if code and str(code).strip():
            _add_heading("代码实现", 1)
            _add_body("以下为本次任务的核心 Python 实现代码：")
            _add_code_block(code)
            stdout = c.get("stdout", "")
            if stdout and str(stdout).strip():
                _add_heading("代码执行输出", 2)
                _add_code_block(stdout)
            doc.add_paragraph()

        # ── 文献列表 ──
        papers = c.get("literature", [])
        if papers:
            _add_heading("参考文献", 1)
            for i, p in enumerate(papers, 1):
                t   = p.get("title", "Unknown")
                aus = p.get("authors", [])
                au  = ", ".join(aus[:3]) + (" et al." if len(aus) > 3 else "")
                yr  = p.get("published", "")[:4]
                url = p.get("url", "")
                ref_p = doc.add_paragraph(style='List Number')
                ref_p.add_run(f"{au} ({yr}). {t}. ").font.size = Pt(11)
                if url:
                    url_run = ref_p.add_run(url)
                    url_run.font.size  = Pt(10)
                    url_run.font.color.rgb = RGBColor(0x0, 0x56, 0xB3)
                # 结构化摘要
                ss = p.get("structured_summary")
                if ss and isinstance(ss, dict):
                    detail_p = doc.add_paragraph()
                    detail_p.paragraph_format.left_indent = Cm(1.0)
                    detail_p.add_run("方法：").bold = True
                    detail_p.add_run(ss.get("method","") + "  ")
                    detail_p.add_run("结论：").bold = True
                    detail_p.add_run(ss.get("conclusion",""))
                    detail_p.runs[-1].font.size = Pt(10)

        # ── 评估信息 ──
        ev = c.get("evaluation", {})
        if ev:
            _add_heading("质量评估", 1)
            score = ev.get("overall_score", "N/A")
            _add_body(f"综合评分：{score}　|　"
                      f"准确性：{ev.get('accuracy','N/A')}　|　"
                      f"完整性：{ev.get('completeness','N/A')}　|　"
                      f"格式规范：{ev.get('format_quality','N/A')}")
            fb = ev.get("feedback", "")
            if fb:
                _add_body(f"评审意见：{fb}")

        doc.save(str(path))

    # ── Markdown（详细版）────────────────────────────────────
    def _to_markdown(self, c: dict) -> str:
        title  = c.get("title", "MindPilot 科研报告")
        query  = c.get("query", "")
        ts_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        lines = [
            f"# {title}",
            "",
            f"> **研究问题**：{query}",
            f">",
            f"> **生成时间**：{ts_str}  |  **系统**：MindPilot 多Agent科研助手 v1.0",
            "",
            "---",
            "",
        ]

        # 摘要
        abstract = c.get("abstract", "")
        if abstract:
            lines += ["## 摘要", "", abstract, ""]

        # 正文章节
        for sec in c.get("sections", []):
            h     = sec.get("heading", "")
            body  = sec.get("body", "")
            level = sec.get("level", 1)
            prefix = "#" * (level + 1)
            if h:
                lines += [f"{prefix} {h}", ""]
            if body:
                lines += [str(body), ""]

        # 代码
        code = c.get("code", "")
        if code and str(code).strip():
            lines += [
                "## 核心代码实现",
                "",
                "```python",
                str(code),
                "```",
                "",
            ]
            stdout = c.get("stdout", "")
            if stdout and str(stdout).strip():
                lines += [
                    "### 执行输出",
                    "",
                    "```",
                    str(stdout)[:2000],
                    "```",
                    "",
                ]

        # 文献
        papers = c.get("literature", [])
        if papers:
            lines += ["## 参考文献", ""]
            for i, p in enumerate(papers, 1):
                t   = p.get("title", "Unknown")
                aus = p.get("authors", [])
                au  = ", ".join(aus[:3]) + (" et al." if len(aus) > 3 else "")
                yr  = p.get("published", "")[:4]
                url = p.get("url", "")
                lines.append(f"{i}. **{t}**  ")
                lines.append(f"   {au} ({yr})  |  [{url}]({url})" if url else f"   {au} ({yr})")
                ss = p.get("structured_summary")
                if ss and isinstance(ss, dict):
                    lines.append(f"   - 方法：{ss.get('method','')}")
                    lines.append(f"   - 结论：{ss.get('conclusion','')}")
                    lines.append(f"   - 局限：{ss.get('limitation','')}")
                lines.append("")

        # 图表
        charts = c.get("charts", [])
        if charts:
            lines += ["## 可视化图表", ""]
            for ch in charts:
                if ch:
                    name = Path(ch).name
                    lines += [f"![{name}]({ch})", ""]

        # 评估
        ev = c.get("evaluation", {})
        if ev:
            score = ev.get("overall_score", "N/A")
            lines += [
                "## 质量评估",
                "",
                f"| 维度 | 得分 |",
                f"|------|------|",
                f"| 综合评分 | {score} |",
                f"| 准确性   | {ev.get('accuracy','N/A')} |",
                f"| 完整性   | {ev.get('completeness','N/A')} |",
                f"| 格式规范 | {ev.get('format_quality','N/A')} |",
                "",
                f"> {ev.get('feedback','')}",
                "",
            ]

        return "\n".join(lines)

    # ── HTML ──────────────────────────────────────────────────
    def _to_html(self, c: dict) -> str:
        md = self._to_markdown(c)
        title = c.get("title", "MindPilot Report")

        body = md
        body = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', body, flags=re.M)
        body = re.sub(r'^### (.+)$',  r'<h3>\1</h3>', body, flags=re.M)
        body = re.sub(r'^## (.+)$',   r'<h2>\1</h2>', body, flags=re.M)
        body = re.sub(r'^# (.+)$',    r'<h1>\1</h1>', body, flags=re.M)
        body = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', body)
        body = re.sub(r'\*(.+?)\*',     r'<em>\1</em>', body)
        body = re.sub(r'```python\n(.*?)```', r'<pre><code class="python">\1</code></pre>', body, flags=re.DOTALL)
        body = re.sub(r'```\n(.*?)```',       r'<pre><code>\1</code></pre>', body, flags=re.DOTALL)
        body = re.sub(r'!\[(.+?)\]\((.+?)\)', r'<img src="\2" alt="\1" style="max-width:100%;border-radius:6px">', body)
        body = re.sub(r'\[(.+?)\]\((.+?)\)',  r'<a href="\2">\1</a>', body)
        body = re.sub(r'^> (.+)$', r'<blockquote>\1</blockquote>', body, flags=re.M)
        body = re.sub(r'^\| (.+)$', r'<tr><td>\1</td></tr>', body, flags=re.M)
        body = re.sub(r'^- (.+)$', r'<li>\1</li>', body, flags=re.M)
        body = re.sub(r'^\d+\. (.+)$', r'<li>\1</li>', body, flags=re.M)
        body = body.replace("---", "<hr>")
        body = re.sub(r'\n\n', '</p><p>', body)
        body = f"<p>{body}</p>"

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
body{{font-family:-apple-system,'Segoe UI',sans-serif;max-width:960px;margin:40px auto;padding:0 24px;color:#222;line-height:1.8}}
h1{{color:#1F4E79;border-bottom:3px solid #2E75B6;padding-bottom:10px;font-size:2em}}
h2{{color:#2E75B6;border-left:4px solid #2E75B6;padding-left:12px;margin-top:2em}}
h3{{color:#444;margin-top:1.5em}}
blockquote{{background:#EFF6FF;border-left:4px solid #2E75B6;margin:0;padding:12px 20px;border-radius:4px}}
pre{{background:#1E1E1E;color:#D4D4D4;padding:18px;border-radius:8px;overflow-x:auto;font-size:13px;line-height:1.5}}
code{{font-family:'Consolas','Courier New',monospace}}
table{{width:100%;border-collapse:collapse;margin:1em 0}}
td,th{{border:1px solid #ddd;padding:8px 12px}}
th{{background:#2E75B6;color:#fff}}
tr:nth-child(even){{background:#F7FBFF}}
img{{border-radius:8px;box-shadow:0 2px 8px rgba(0,0,0,.12);max-width:100%}}
.footer{{text-align:center;color:#999;font-size:12px;margin-top:60px;padding-top:20px;border-top:1px solid #eee}}
</style>
</head>
<body>
{body}
<div class="footer">Generated by MindPilot v1.0 &nbsp;|&nbsp; {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>
</body>
</html>"""
